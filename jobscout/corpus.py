"""Reading the applications you have already written.

The expected shape is one folder per company::

    <applications_dir>/
        Acme Corp/
            resume_acme.pdf
            cover_letter.pdf
            acme_jd.txt
        Globex/
            ...

but a flat folder of documents works too — anything it cannot attribute to a
company is still read for profile-building.

Nothing here leaves your machine except as prompt text sent to the model you
configured, and only the excerpts the caller asks for.
"""
from __future__ import annotations

import datetime as dt
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Set

TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".text", ".csv", ".json"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}
SKIP_NAMES = {".DS_Store", "Thumbs.db", "Icon\r"}

#: Filename hints, checked in order — the first match wins.
KIND_PATTERNS = (
    ("cover_letter", re.compile(r"cover[\W_]*letter|coverletter|\bcl\b", re.I)),
    ("resume", re.compile(r"resume|\bcv\b|curriculum[\W_]*vitae", re.I)),
    ("job_description", re.compile(r"\bjd\b|job[\W_]*desc|posting|requisition|role", re.I)),
    ("correspondence", re.compile(r"recruiter|thread|email|interview|call|prep|screen", re.I)),
)

#: Corporate suffixes stripped when comparing two company names.
_COMPANY_NOISE = re.compile(
    r"\b(inc|inc\.|llc|l\.l\.c|ltd|limited|corp|corporation|co|company|"
    r"holdings|group|technologies|technology|labs|laboratory|laboratories|"
    r"national|the)\b",
    re.I,
)


def normalize_company(name: str) -> str:
    """A comparison key for a company name.

    ``"Sandia National Laboratories"`` and ``"Sandia Labs"`` both reduce to
    ``"sandia"``, so an application you already sent is recognised when a job
    board spells the employer differently.
    """
    text = re.sub(r"[^\w\s&-]", " ", name or "")
    text = _COMPANY_NOISE.sub(" ", text)
    text = re.sub(r"[\s_-]+", " ", text).strip().lower()
    return text


@dataclass
class Document:
    path: Path
    company: str
    kind: str
    text: str
    modified: dt.date
    error: str = ""

    @property
    def name(self) -> str:
        return self.path.name

    def excerpt(self, limit: int) -> str:
        if len(self.text) <= limit:
            return self.text
        return self.text[:limit].rstrip() + "\n…[truncated]"


@dataclass
class Application:
    """Everything found under one company folder."""

    company: str
    documents: List[Document] = field(default_factory=list)

    @property
    def key(self) -> str:
        return normalize_company(self.company)

    @property
    def last_touched(self) -> Optional[dt.date]:
        dates = [d.modified for d in self.documents if d.modified]
        return max(dates) if dates else None

    def kinds(self) -> Set[str]:
        return {d.kind for d in self.documents}


@dataclass
class Corpus:
    root: Path
    applications: List[Application] = field(default_factory=list)
    loose: List[Document] = field(default_factory=list)
    skipped: List[str] = field(default_factory=list)

    @property
    def documents(self) -> List[Document]:
        docs: List[Document] = []
        for app in self.applications:
            docs.extend(app.documents)
        docs.extend(self.loose)
        return docs

    def company_keys(self) -> Set[str]:
        return {app.key for app in self.applications if app.key}

    def company_names(self) -> List[str]:
        return [app.company for app in self.applications]


# --- text extraction -------------------------------------------------------

def _read_pdf(path: Path) -> str:
    import logging

    from pypdf import PdfReader  # imported lazily so `--help` needs no deps

    # pypdf narrates every malformed cross-reference in exported PDFs; the text
    # comes out fine and the noise would drown the CLI's own output.
    logging.getLogger("pypdf").setLevel(logging.ERROR)

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # a single broken page should not lose the file
            pages.append("[unreadable page: %s]" % exc)
    return "\n".join(pages)


def _read_docx(path: Path) -> str:
    import docx  # lazy, as above

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in PDF_SUFFIXES:
        return _read_pdf(path)
    if suffix in DOCX_SUFFIXES:
        return _read_docx(path)
    if suffix in TEXT_SUFFIXES:
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError("unsupported file type %s" % (suffix or "(none)"))


def classify(path: Path) -> str:
    stem = path.stem
    for kind, pattern in KIND_PATTERNS:
        if pattern.search(stem):
            return kind
    if path.suffix.lower() in PDF_SUFFIXES:
        return "document"
    return "notes"


def _load_document(path: Path, company: str) -> Optional[Document]:
    if path.name in SKIP_NAMES or path.name.startswith("."):
        return None
    suffix = path.suffix.lower()
    if suffix not in TEXT_SUFFIXES | PDF_SUFFIXES | DOCX_SUFFIXES:
        return None
    try:
        modified = dt.date.fromtimestamp(path.stat().st_mtime)
    except OSError:
        modified = dt.date.today()
    try:
        text = extract_text(path)
        error = ""
    except Exception as exc:
        text = ""
        error = str(exc)
    return Document(path=path, company=company, kind=classify(path),
                    text=text.strip(), modified=modified, error=error)


def load_corpus(root: Path) -> Corpus:
    """Walk ``root`` and read every document it understands."""
    corpus = Corpus(root=root)
    if not root.is_dir():
        raise FileNotFoundError("applications folder not found: %s" % root)

    for entry in sorted(root.iterdir()):
        if entry.name in SKIP_NAMES or entry.name.startswith("."):
            continue
        if entry.is_dir():
            app = Application(company=entry.name)
            for child in sorted(entry.rglob("*")):
                if not child.is_file():
                    continue
                doc = _load_document(child, entry.name)
                if doc is None:
                    corpus.skipped.append(str(child.relative_to(root)))
                else:
                    app.documents.append(doc)
            if app.documents:
                corpus.applications.append(app)
        elif entry.is_file():
            doc = _load_document(entry, "")
            if doc is None:
                corpus.skipped.append(entry.name)
            else:
                corpus.loose.append(doc)
    return corpus


def summarize(corpus: Corpus) -> str:
    """A short, printable inventory — filenames only, no document content."""
    lines = ["%d application folder(s), %d document(s)"
             % (len(corpus.applications), len(corpus.documents))]
    for app in corpus.applications:
        touched = app.last_touched.isoformat() if app.last_touched else "unknown"
        lines.append("  %-28s %2d docs  last touched %s  [%s]"
                     % (app.company, len(app.documents), touched,
                        ", ".join(sorted(app.kinds()))))
    if corpus.loose:
        lines.append("  (%d loose document(s) not attributed to a company)" % len(corpus.loose))
    unreadable = [d for d in corpus.documents if d.error]
    if unreadable:
        lines.append("  %d document(s) could not be read" % len(unreadable))
    return "\n".join(lines)
